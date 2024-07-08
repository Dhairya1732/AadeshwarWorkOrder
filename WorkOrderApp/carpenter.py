from docx import Document
import os
from docx.shared import Pt
from docx.enum.text import WD_UNDERLINE, WD_ALIGN_PARAGRAPH
from utils import set_cell_border, set_run_font
from docx2pdf import convert 

class CarpenterWorkOrder:
    def __init__(self,backend):
        self.backend = backend

    def create_carpenter_order(self, order_data, template_path, pdf_filename):
        pdf_output_path = os.path.join(self.backend.download_path, pdf_filename)     
        try:
            docx_filename = pdf_filename.replace('.pdf', '.docx')
            docx_output_path = os.path.join(self.backend.download_path, docx_filename)
            self.process_order_data(order_data, template_path, docx_output_path)
            convert(docx_output_path, pdf_output_path)
            os.remove(docx_output_path)
            print(f"Successfully processed carpenter order for {order_data[0]['To be shipped Before']}.")
        except Exception as e:
            print(f"Failed to process carpenter order: {e}")

    def process_order_data(self, order_data, template_path, output_path):
        try:
            doc = Document(template_path)
        except Exception as e:
            print(f"Error opening template: {e}")
            return
        
        table = doc.tables[0]  

        if len(table.rows) >= 2 and len(table.rows[1].cells) > 0:
            cell = table.rows[1].cells[0]
            cell.text = cell.text.replace('[CarpenterName]',order_data[0]['Carpenter Team'])
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    set_run_font(run, 16)
                    run.font.underline = WD_UNDERLINE.SINGLE
                    

        if len(table.rows) >= 3 and len(table.rows[2].cells) > 0:
            cell = table.rows[2].cells[0]
            cell.text = cell.text.replace('[Delivery Date]', order_data[0]['To be shipped Before'])
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    set_run_font(run, 16)
                    
        for order_data in order_data:
            row = table.add_row()
            cells = row.cells
            cells[0].text = str(order_data['QTY'])
            cells[1].text = order_data['OrderNo']
            cells[2].text = order_data['To be shipped Before']
            cells[3].text = order_data['Your SKU ID']
            cells[4].text = str(order_data['Carpenter Inches'])
            cells[5].text = str(order_data['TotalInches'])

            for cell in cells:
                set_cell_border(cell)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run, 11)
        
        try:
            doc.save(output_path)
            print(f"Document saved to {output_path}")
        except Exception as e:
            print(f"Error saving document: {e}")