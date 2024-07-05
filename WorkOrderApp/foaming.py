from docx import Document
from docx.shared import Pt
import os
from docx2pdf import convert 
from utils import set_run_font
import requests
from io import BytesIO
from work_order_backend import WorkOrderAppBackend

backend = WorkOrderAppBackend()

class FoamingWorkOrder:
    def create_work_order(self, order_data, template_path, pdf_filename):
        pdf_output_path = os.path.join(backend.download_path, pdf_filename)     
        try:
            docx_filename = pdf_filename.replace('.pdf', '.docx')
            docx_output_path = os.path.join(backend.download_path, docx_filename)
            self.process_work_order(order_data, template_path, docx_output_path, image_width=Pt(100), image_height=Pt(100))
            convert(docx_output_path, pdf_output_path)
            print(f"Successfully processed work order {order_data['OrderNo']}.")
            os.remove(docx_output_path)
        except Exception as e:
            print(f"Failed to process work order {order_data['OrderNo']}: {e}")

    def process_work_order(self, data, template_path, output_path, image_width=None, image_height=None):
        try:
            doc = Document(template_path)
        except Exception as e:
            print(f"Error opening template: {e}")
            return
        table = doc.tables[0]
        for key, value in data.items():
            if isinstance(value, str):
                value = value.strip()
            found = False
            for row in table.rows:
                if found:
                    break
                for cell in row.cells:
                    if found:
                        break
                    for paragraph in cell.paragraphs:
                        if f'[{key}]' in paragraph.text:    
                            if key == 'Image url':
                                paragraph.text = paragraph.text.replace(f'[{key}]', '')
                                self.insert_image_from_url(cell, value, width=image_width, height=image_height)
                            else:
                                paragraph.text = paragraph.text.replace(f'[{key}]', str(value))
                                for run in paragraph.runs:
                                    set_run_font(run)
                                if template_path==backend.carpenter_template_path and key=='Order Confirmed Date':
                                    run.font.size = Pt(16)
                            found = True
                            break        

        try:
            doc.save(output_path)
            print(f"Document saved to {output_path}")
        except Exception as e:
            print(f"Error saving document: {e}")

    def insert_image_from_url(self, cell, image_url, width=None, height=None):
        try:
            response = requests.get(image_url)
            image_stream = BytesIO(response.content)
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(image_stream, width=width, height=height)
        except Exception as e:
            print(f"Error inserting image from URL {image_url}: {e}")