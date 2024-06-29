from docx import Document
from docx.shared import Pt
import pandas as pd
import requests
from io import BytesIO
from PyQt5.QtCore import QSettings, QDir
from utils import set_run_font
from docx2pdf import convert 
import os
from datetime import timedelta, datetime

COMPANY_NAME = "Aadeshwar"
APP_NAME = "WorkOrderGenerator"

class WorkOrderAppBackend:
    def __init__(self):
        self.csv_path = ''
        self.foaming_template_path = ''
        self.carpenter_template_path = ''
        self.sales_template_path = ''
        self.download_path = QDir.homePath()  
        self.settings = QSettings(COMPANY_NAME, APP_NAME)
        self.load_settings()

    def load_settings(self):
        self.download_path = self.settings.value("download_path", QDir.homePath())
        self.foaming_template_path = self.settings.value("foaming_template_path", '')
        self.carpenter_template_path = self.settings.value("carpenter_template_path", '')
        self.sales_template_path = self.settings.value("sales_template_path", '')


    def save_settings(self):
        self.settings.setValue("download_path", self.download_path)

    def set_csv_path(self, path):
        self.csv_path = path
    
    def set_foaming_template_path(self, path):
        self.foaming_template_path = path
        self.settings.setValue("foaming_template_path", self.foaming_template_path)

    def set_carpenter_template_path(self, path):
        self.carpenter_template_path = path
        self.settings.setValue("carpenter_template_path", self.carpenter_template_path)

    def set_sales_template_path(self, path):
        self.sales_template_path = path
        self.settings.setValue("sales_template_path", self.sales_template_path)

    def set_download_path(self, path):
        self.download_path = path
        self.save_settings()

    def generate_work_order(self):
        orders = pd.read_csv(self.csv_path)
        fabric_sheet = pd.read_excel(r"D:\AadeshwarWorkOrder\Database\Fabric_Work Order revise new.xlsx",sheet_name="Fabric")
        #inches = pd.read_excel("D:\AadeshwarWorkOrder\Database\Daily Update sheet.xlsx",sheet_name="Sheet4")
        orders = orders.tail(1)
        columns_to_ignore = ['Unit Price', 'TOTAL', 'Shipping Address', 'status', 'Promised Delivery Date']
        
        for index, row in orders.iterrows():
            order_data = row.to_dict()
            for column in columns_to_ignore:
                order_data.pop(column, None)

            relevant_columns = ['SKU_ID', 'Legs', 'Legs Quantity', 'Cushion Qty', 'Cushion Fabric', 'Sofa Fabric', 'Legs Finish', 'Legs Assembly', 'Cushions', 'Cushion Size']
            fabric_sheet = fabric_sheet[relevant_columns]
            fabric_sheet = fabric_sheet.fillna("")
            sku_id = order_data['SKU ID'] 
            # Check if the 'Fabric' sheet contains the 'SKU_ID' column
            if 'SKU_ID' in fabric_sheet.columns:
                # Iterate over the orders
                for index, row in orders.iterrows():
                    order_data = row.to_dict()
                    for column in columns_to_ignore:
                        order_data.pop(column, None)

            if 'Order Confirmed Date' in order_data:
                order_data['Order Confirmed Date'] = pd.to_datetime(order_data['Order Confirmed Date'],dayfirst=True).date()

            if 'To be shipped Before' in order_data:
                delivery_date = pd.to_datetime(order_data['To be shipped Before'],dayfirst=True)
                if pd.notnull(delivery_date):
                    adjusted_date = (delivery_date - timedelta(days=2)).date()
                    order_data['To be shipped Before'] = adjusted_date.strftime(f"%d-%m-%Y")

            current_month = datetime.now().strftime("%B")
            order_no = f"G1/{current_month}/{self.current_order_no}"
            order_data['OrderNo'] = order_no

            sku_id = order_data['SKU ID']
            if fabric_sheet['SKU_ID'].isin([sku_id]).any():
                additional_data = fabric_sheet[fabric_sheet['SKU_ID'] == sku_id].to_dict('records')[0]
            else:
                additional_data= {col: "" for col in relevant_columns if col!="SKU_ID"}
            order_data.update(additional_data)

            # if inches['Sheet4'].isin([sku_id]).any():
            #     additional_data = inches['Sheet4'].loc[inches['Sheet4'] == sku_id].to_dict('records')[0]
            #     order_data.update(additional_data)

            self.create_work_order(order_data, self.foaming_template_path, f"foaming_{index+1}.pdf")
            self.create_work_order(order_data, self.carpenter_template_path, f"carpenter_{index+1}.pdf")
            self.create_work_order(order_data, self.sales_template_path, f"sales_{index+1}.pdf")
            self.current_order_no += 1
    
    def create_work_order(self, order_data, template_path, pdf_filename):
        pdf_output_path = os.path.join(self.download_path, pdf_filename)     
        try:
            docx_filename = pdf_filename.replace('.pdf', '.docx')
            docx_output_path = os.path.join(self.download_path, docx_filename)
            self.process_work_order(order_data, template_path, docx_output_path, image_width=Pt(100), image_height=Pt(100))
            convert(docx_output_path, pdf_output_path)
            print(f"Successfully processed work order {order_data['OrderNo']}.")
            os.remove(docx_output_path)
        except Exception as e:
            print(f"Failed to process work order {order_data['OrderNo']}: {e}")

    def process_work_order(self, data, template_path, output_path, image_width=None, image_height=None):
        try:
            doc = Document(template_path)
        except Exception as e:
            print(f"Error opening template: {e}")
            return
        
        for key, value in data.items():
            if isinstance(value, str):
                value = value.strip()
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if f'[{key}]' in paragraph.text:
                                if key == 'Image url':
                                    paragraph.text = paragraph.text.replace(f'[{key}]', '')
                                    self.insert_image_from_url(cell, value, width=image_width, height=image_height)
                                else:
                                    paragraph.text = paragraph.text.replace(f'[{key}]', str(value))
                                    for run in paragraph.runs:
                                        set_run_font(run)
                                    if template_path==self.carpenter_template_path and key=='Order Confirmed Date':
                                        run.font.size = Pt(16)

        try:
            doc.save(output_path)
            print(f"Document saved to {output_path}")
        except Exception as e:
            print(f"Error saving document: {e}")

    def insert_image_from_url(self, cell, image_url, width=None, height=None):
        try:
            response = requests.get(image_url)
            image_stream = BytesIO(response.content)
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(image_stream, width=width, height=height)
        except Exception as e:
            print(f"Error inserting image from URL {image_url}: {e}")