import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
import requests
from io import BytesIO

def set_run_font(run, font_name='Times New Roman', font_size=9):
    run.font.name = font_name
    run.font.size = Pt(font_size)

def insert_image_from_url(cell, image_url):
    try:
        response = requests.get(image_url)
        response.raise_for_status()  # Raise an error for bad responses
        image_stream = BytesIO(response.content)
        
        # Add the image directly to the cell
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(image_stream, width=Inches(1.5))  # Adjust width as needed
    except requests.exceptions.RequestException as e:
        print(f"Failed to insert image from {image_url}: {e}")

def generate_work_order(data, template_path, output_path):
    # Load the template document
    doc = Document(template_path)
    
    # Replace placeholders with actual values
    for key, value in data.items():
        if isinstance(value, str):  # Ensure value is a string before processing
            value = value.strip()  # Remove any leading/trailing whitespace
        
        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Process each paragraph in the cell once
                    for paragraph in cell.paragraphs:
                        if f'[{key}]' in paragraph.text:
                            if key == 'Image url':
                                paragraph.text = paragraph.text.replace(f'[{key}]', '')  # Clear placeholder text
                                insert_image_from_url(cell, value)
                            else:
                                paragraph.text = paragraph.text.replace(f'[{key}]', str(value))
                                for run in paragraph.runs:
                                    set_run_font(run)
    
    # Save the document
    doc.save(output_path)

def main():
    # Load the Excel file
    excel_file = r"D:\AadeshwarWorkOrder\Pendingorders-2024-06-12.xlsx"
    orders = pd.read_excel(excel_file)
    
    # Path to the template document
    template_path = r"D:\AadeshwarWorkOrder\FoamingTemplate.docx"

    # For testing, uncomment the next line to process only the first row
    orders = orders.head(1)
    
    # List of columns to ignore
    columns_to_ignore = ['Unit Price', 'TOTAL', 'SKU ID', 'Shipping Address', 'status', 'Promised Delivery Date']

    # Generate work orders for each row in the Excel file
    for index, row in orders.iterrows():
        order_data = row.to_dict()
        
        # Remove unwanted columns
        for column in columns_to_ignore:
            order_data.pop(column, None)
        
        output_path = f"Work_Order_{index + 1}.docx"  # Using index + 1 for file name
        generate_work_order(order_data, template_path, output_path)
        print(f"Generated work order for Row No.: {index + 1}")

if __name__ == '__main__':
    main()