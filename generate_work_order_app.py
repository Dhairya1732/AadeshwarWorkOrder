import sys
import os
import pandas as pd
from PyQt5.QtWidgets import QApplication, QWidget, QVBoxLayout, QPushButton, QFileDialog, QLabel, QDesktopWidget
from PyQt5.QtGui import QIcon, QFont
from PyQt5.QtCore import QSettings
from docx import Document
from docx.shared import Pt
import requests
from io import BytesIO
from pyshortcuts import make_shortcut

class WorkOrderApp(QWidget):
    def __init__(self):
        super().__init__()
        self.initUI()

    def initUI(self):
        layout = QVBoxLayout()
        
        self.excel_btn = QPushButton('Upload Excel File')
        self.excel_btn.clicked.connect(self.upload_excel)
        self.set_button_font_size(self.excel_btn, 12)
        self.excel_btn.setStyleSheet('QPushButton { padding: 15px; }') 
        layout.addWidget(self.excel_btn)
        
        self.template_btn = QPushButton('Upload Word Template')
        self.template_btn.clicked.connect(self.upload_template)
        self.set_button_font_size(self.template_btn, 12)
        self.template_btn.setStyleSheet('QPushButton { padding: 15px; }') 
        layout.addWidget(self.template_btn)
        
        self.generate_btn = QPushButton('Generate Work Order')
        self.generate_btn.clicked.connect(self.generate_work_order)
        self.set_button_font_size(self.generate_btn, 12)
        self.generate_btn.setStyleSheet('QPushButton { padding: 15px; }') 
        layout.addWidget(self.generate_btn)
        
        self.status_label = QLabel('')
        layout.addWidget(self.status_label)
        
        self.setLayout(layout)
        self.setWindowTitle('Work Order Generator')
        self.setWindowIcon(QIcon('app_icon.ico'))  
        self.excel_path = ''
        self.template_path = ''

        self.settings = QSettings("MyCompany", "WorkOrderGenerator")

        template_path = self.settings.value("template_path")
        if template_path and os.path.isfile(template_path):
            self.template_path = template_path
            self.status_label.setText('Word Template Loaded')
        else:
            self.status_label.setText('Please upload Word Template')

        # Adjust window size and center on the screen
        self.resize(600, 500) 
        self.center()  

    def center(self):
        # Calculate center position based on screen geometry
        qr = self.frameGeometry()
        cp = QDesktopWidget().availableGeometry().center()
        qr.moveCenter(cp)
        self.move(qr.topLeft())

    def set_button_font_size(self, button, font_size):
        font = QFont()
        font.setPointSize(font_size)
        button.setFont(font)
    
    def upload_excel(self):
        options = QFileDialog.Options()
        options |= QFileDialog.ReadOnly
        file_path, _ = QFileDialog.getOpenFileName(self, 'Upload Excel File', '', 'Excel Files (*.xlsx)', options=options)
        if file_path:
            self.excel_path = file_path
            self.status_label.setText('Excel File Uploaded')
    
    def upload_template(self):
        options = QFileDialog.Options()
        options |= QFileDialog.ReadOnly
        file_path, _ = QFileDialog.getOpenFileName(self, 'Upload Word Template', '', 'Word Files (*.docx)', options=options)
        if file_path:
            self.template_path = file_path
            self.settings.setValue("template_path", self.template_path)
            self.status_label.setText('Word Template Uploaded')
    
    def generate_work_order(self):
        if not self.excel_path or not self.template_path:
            self.status_label.setText('Please upload both files')
            return

        orders = pd.read_excel(self.excel_path)
        orders = orders.head(1)
        columns_to_ignore = ['Unit Price', 'TOTAL', 'SKU ID', 'Shipping Address', 'status', 'Promised Delivery Date']
        
        for index, row in orders.iterrows():
            order_data = row.to_dict()
            for column in columns_to_ignore:
                order_data.pop(column, None)
            
            output_path = f"Work_Order_{index + 1}.docx"
            self.process_work_order(order_data, self.template_path, output_path, image_width=Pt(100), image_height=Pt(100))
        
        self.status_label.setText('Work Order Generated')
    
    def process_work_order(self, data, template_path, output_path, image_width=None, image_height=None):
        doc = Document(template_path)
        for key, value in data.items():
            if isinstance(value, str):
                value = value.strip()
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if f'[{key}]' in paragraph.text:
                                if key == 'Image url':
                                    paragraph.text = paragraph.text.replace(f'[{key}]', '')
                                    self.insert_image_from_url(cell, value, width=image_width, height=image_height)
                                else:
                                    paragraph.text = paragraph.text.replace(f'[{key}]', str(value))
                                    for run in paragraph.runs:
                                        self.set_run_font(run)
        doc.save(output_path)

    def set_run_font(self, run, font_name='Times New Roman', font_size=9):
        run.font.name = font_name
        run.font.size = Pt(font_size)
    
    def insert_image_from_url(self, cell, image_url, width=None, height=None):
        response = requests.get(image_url)
        image_stream = BytesIO(response.content)
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(image_stream, width=width, height=height)

def create_shortcut():
    script_path = os.path.abspath('dist/work_order_generator.exe')
    make_shortcut(script_path, name='WorkOrderGenerator', description='Work Order Generator', icon='app_icon.ico')

if __name__ == '__main__':
    app = QApplication(sys.argv)
    
    ex = WorkOrderApp()
    ex.show()

    # Create the desktop shortcut
    create_shortcut()
    
    sys.exit(app.exec_())