from docx import Document
import os
from docx2pdf import convert 
from utils import set_cell_border, set_run_font
from docx.enum.text import WD_ALIGN_PARAGRAPH

class SalesSummary:
    def __init__(self, backend):
        self.backend = backend

    def create_sales_summary(self, orders_data, template_path, pdf_filename):
        pdf_output_path = os.path.join(self.backend.download_path, pdf_filename)
        try:
            docx_filename = pdf_filename.replace('.pdf', '.docx')
            docx_output_path = os.path.join(self.backend.download_path, docx_filename)
            self.process_sales_summary(orders_data, template_path, docx_output_path)
            convert(docx_output_path, pdf_output_path)
            os.remove(docx_output_path)
            print(f"Successfully processed sales summary for {orders_data[0]['To be shipped Before']}.")
        except Exception as e:
            print(f"Failed to process sales summary: {e}")

    def process_sales_summary(self, orders_data, template_path, output_path):
        try:
            doc = Document(template_path)
        except Exception as e:
            print(f"Error opening template: {e}")
            return

        table = doc.tables[0]
        sr_no = 1
        sales_no = self.backend.current_sales_no

        if len(table.rows) >= 1 and len(table.rows[0].cells) > 0:
            cell = table.rows[0].cells[0]
            cell.text = cell.text.replace('[SalesNo]', str(sales_no))
            for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        set_run_font(run, 16)

        if len(table.rows) >= 2 and len(table.rows[1].cells) > 0:
            cell = table.rows[1].cells[0]
            cell.text = cell.text.replace('[Delivery Date]', orders_data[0]['To be shipped Before'])
            for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run, 11)

        for order_data in orders_data:
            row = table.add_row()
            cells = row.cells
            cells[0].text = str(sr_no)
            cells[1].text = order_data['OrderNo']
            cells[2].text = order_data['Customer Name']
            cells[3].text = order_data['Your SKU ID']
            cells[4].text = order_data['Sofa Fabric']
            cells[5].text = str(order_data['QTY'])

            for cell in cells:
                set_cell_border(cell)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run, 11)

            for i in range(len(cells), 10):  
                cell = row.add_cell()
                set_cell_border(cell)

            sr_no += 1
            
        self.backend.current_sales_no += 1
        
        try:
            doc.save(output_path)
            print(f"Sales summary document saved to {output_path}")
        except Exception as e:
            print(f"Error saving document: {e}")